from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
SCRIPT_DIR = ROOT / "scripts" / "cfo-kai"


def load_module(name: str):
    spec = importlib.util.spec_from_file_location(name, SCRIPT_DIR / f"{name}.py")
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def sample_finance_data() -> dict:
    return {
        "generated_at": "2026-06-05T09:30:00Z",
        "as_of": "Finance snapshot test · 2026-06-05",
        "currency": "EUR",
        "cash_on_hand_eur": 133039,
        "runway_months": 12,
        "break_even_label": "M6 · Jan 2027",
        "monthly_burn": {"actual_eur": 2068, "plan_eur": 2068, "delta_eur": 0},
        "cost_lines": [
            {"label": "API", "amount_eur": 400, "fixed": False, "paid_by": "Company", "note": "usage"},
            {"label": "Claude Max", "amount_eur": 220, "fixed": True, "paid_by": "Felix", "note": "2 seats"},
        ],
        "paid_by": [
            {"name": "Company", "value_eur": 400},
            {"name": "Felix", "value_eur": 220},
        ],
        "forecast_6m": [
            {"month": "Jun", "cash_eur": 133039, "burn_eur": 2068},
            {"month": "Jul", "cash_eur": 130971, "burn_eur": 2068},
        ],
        "pilot_health": [
            {"name": "Pilot funnel", "status": "green", "note": "Plan intact"},
        ],
    }


def rgb_endswith(cell, suffix: str) -> bool:
    color = cell.font.color.rgb if cell.font and cell.font.color and cell.font.color.type == "rgb" else ""
    return bool(color and color.upper().endswith(suffix))


def fill_endswith(cell, suffix: str) -> bool:
    color = cell.fill.fgColor.rgb if cell.fill and cell.fill.fgColor and cell.fill.fgColor.type == "rgb" else ""
    return bool(color and color.upper().endswith(suffix))


def test_export_finance_xlsx_uses_financial_model_conventions(tmp_path):
    export_finance_xlsx = load_module("export_finance_xlsx")

    path = export_finance_xlsx.build_workbook(
        sample_finance_data(),
        output_dir=tmp_path,
        recalc=False,
        now_label="2026-06",
    )

    assert path.name == "cfo-kai-finance-2026-06.xlsx"
    wb = load_workbook(path, data_only=False)
    assert wb.sheetnames == ["Summary", "Inputs", "Costs", "Forecast", "Assumptions"]

    summary = wb["Summary"]
    inputs = wb["Inputs"]
    costs = wb["Costs"]
    forecast = wb["Forecast"]

    assert summary["B4"].value == "=Inputs!B4"
    assert summary["B8"].value == "=B6-B7"
    assert summary["B9"].value == "=SUM(Costs!B2:B3)"
    assert forecast["D3"].value == "=B3-B2"

    assert rgb_endswith(inputs["B4"], "0000FF")
    assert rgb_endswith(summary["B4"], "000000")
    assert fill_endswith(inputs["B8"], "FFFF00")
    assert "€" in costs["B2"].number_format

    assert export_finance_xlsx.scan_formula_errors(path) == {}


def test_report_monthly_docx_matches_finance_snapshot(tmp_path):
    report_monthly = load_module("report_monthly")

    path = report_monthly.build_report(
        sample_finance_data(),
        output_dir=tmp_path,
        output_format="docx",
        now_label="2026-06",
    )

    assert path.name == "cfo-kai-board-report-2026-06.docx"
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "CFO-Kai Monthly Board Report" in text
    assert "Finance snapshot test · 2026-06-05" in text
    assert "Break-even: M6 · Jan 2027" in text
    assert "Source: GET /api/finance" in text

    kpi_rows = [[cell.text for cell in row.cells] for row in doc.tables[0].rows]
    assert ["Cash on hand", "133,039 EUR"] in kpi_rows
    assert ["Monthly burn", "2,068 EUR"] in kpi_rows
    assert ["Runway", "12.0 months"] in kpi_rows
