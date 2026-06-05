#!/usr/bin/env python3
"""Build a CFO-Kai monthly board report from a FinanceData snapshot."""
from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parent))

from cfo_common import (  # noqa: E402
    DEFAULT_BASE,
    drive_link_for,
    format_eur,
    format_months,
    load_or_fetch,
    resolve_output_dir,
    safe_month_label,
)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - exercised on misconfigured bridge hosts.
    raise SystemExit("ERROR: python-docx is required for report_monthly.py") from exc


ACCENT = RGBColor(31, 78, 121)
LIGHT = RGBColor(91, 115, 132)


def build_report(
    data: dict[str, Any],
    output_dir: str | Path | None = None,
    output_format: str = "auto",
    now_label: str | None = None,
) -> Path:
    month = safe_month_label(now_label)
    out_dir = resolve_output_dir(output_dir)
    fmt = resolve_format(output_format)
    if fmt == "pdf":
        path = out_dir / f"cfo-kai-board-report-{month}.pdf"
        build_pdf_report(data, path)
        return path
    path = out_dir / f"cfo-kai-board-report-{month}.docx"
    build_docx_report(data, path)
    return path


def resolve_format(value: str) -> str:
    if value not in {"auto", "pdf", "docx"}:
        raise ValueError("output format must be auto, pdf, or docx")
    if value == "auto":
        return "pdf" if reportlab_available() else "docx"
    if value == "pdf" and not reportlab_available():
        raise RuntimeError("reportlab is required for --format pdf; use --format docx on this host")
    return value


def reportlab_available() -> bool:
    try:
        import reportlab  # noqa: F401
        return True
    except ImportError:
        return False


def build_docx_report(data: dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("CFO-Kai Monthly Board Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.add_run(data["as_of"]).italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Finance snapshot is consistent with GET /api/finance. "
        f"Cash is {format_eur(data['cash_on_hand_eur'])}, burn is "
        f"{format_eur(data['monthly_burn']['actual_eur'])} per month, runway is "
        f"{format_months(float(data['runway_months']))}, and break-even is tracked as "
        f"{data['break_even_label']}."
    )
    doc.add_paragraph(f"Break-even: {data['break_even_label']}")

    doc.add_heading("KPI Table", level=1)
    kpi_table = doc.add_table(rows=1, cols=2)
    kpi_table.style = "Table Grid"
    kpi_table.rows[0].cells[0].text = "KPI"
    kpi_table.rows[0].cells[1].text = "Value"
    for label, value in [
        ("Cash on hand", format_eur(data["cash_on_hand_eur"])),
        ("Monthly burn", format_eur(data["monthly_burn"]["actual_eur"])),
        ("Plan burn", format_eur(data["monthly_burn"]["plan_eur"])),
        ("Burn delta", format_eur(data["monthly_burn"]["delta_eur"])),
        ("Runway", format_months(float(data["runway_months"]))),
        ("Break-even", str(data["break_even_label"])),
    ]:
        cells = kpi_table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    doc.add_heading("Cost Lines", level=1)
    cost_table = doc.add_table(rows=1, cols=5)
    cost_table.style = "Table Grid"
    for idx, header in enumerate(["Cost", "Amount", "Fixed", "Paid by", "Note"]):
        cost_table.rows[0].cells[idx].text = header
    for line in data["cost_lines"]:
        cells = cost_table.add_row().cells
        cells[0].text = str(line.get("label", ""))
        cells[1].text = format_eur(line.get("amount_eur", 0))
        cells[2].text = "yes" if line.get("fixed") else "no"
        cells[3].text = str(line.get("paid_by", ""))
        cells[4].text = str(line.get("note", ""))

    doc.add_heading("Forecast", level=1)
    forecast_table = doc.add_table(rows=1, cols=3)
    forecast_table.style = "Table Grid"
    for idx, header in enumerate(["Month", "Cash", "Burn"]):
        forecast_table.rows[0].cells[idx].text = header
    for point in data["forecast_6m"]:
        cells = forecast_table.add_row().cells
        cells[0].text = str(point.get("month", ""))
        cells[1].text = format_eur(point.get("cash_eur", 0))
        cells[2].text = format_eur(point.get("burn_eur", 0))

    doc.add_heading("Assumptions", level=1)
    for row in data["pilot_health"]:
        doc.add_paragraph(
            f"{row.get('name', '')}: {row.get('status', '')} - {row.get('note', '')}",
            style="List Bullet",
        )

    doc.add_heading("Sources", level=1)
    doc.add_paragraph("Source: GET /api/finance")
    doc.add_paragraph(f"Snapshot as_of: {data['as_of']}")
    doc.add_paragraph(f"Snapshot generated_at: {data.get('generated_at', 'not supplied')}")

    style_report_tables(doc)
    doc.save(path)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for level in (1, 2):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = ACCENT


def style_report_tables(doc: Document) -> None:
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)
                    paragraph.paragraph_format.space_after = Pt(2)
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = ACCENT


def build_pdf_report(data: dict[str, Any], path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    story: list[Any] = []
    story.append(Paragraph(_pdf_text("CFO-Kai Monthly Board Report"), styles["Title"]))
    story.append(Paragraph(_pdf_text(data["as_of"]), styles["Italic"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Executive Summary", styles["Heading1"]))
    story.append(Paragraph(_pdf_text(
        "Finance snapshot is consistent with GET /api/finance. "
        f"Cash is {format_eur(data['cash_on_hand_eur'])}, burn is "
        f"{format_eur(data['monthly_burn']['actual_eur'])} per month, runway is "
        f"{format_months(float(data['runway_months']))}, and break-even is "
        f"{data['break_even_label']}."
    ), styles["BodyText"]))
    story.append(Spacer(1, 12))

    kpi_rows = [["KPI", "Value"]] + [
        ["Cash on hand", format_eur(data["cash_on_hand_eur"])],
        ["Monthly burn", format_eur(data["monthly_burn"]["actual_eur"])],
        ["Plan burn", format_eur(data["monthly_burn"]["plan_eur"])],
        ["Burn delta", format_eur(data["monthly_burn"]["delta_eur"])],
        ["Runway", format_months(float(data["runway_months"]))],
        ["Break-even", _pdf_text(str(data["break_even_label"]))],
    ]
    story.append(_pdf_table(kpi_rows))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Sources", styles["Heading1"]))
    story.append(Paragraph(_pdf_text(f"Source: GET /api/finance; as_of: {data['as_of']}"), styles["BodyText"]))
    doc.build(story)


def _pdf_table(rows: list[list[str]]) -> Any:
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    table = Table([[_pdf_text(str(cell)) for cell in row] for row in rows], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C9D6")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def _pdf_text(value: str) -> str:
    return value.encode("latin-1", "replace").decode("latin-1")


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Build CFO-Kai monthly board report from /api/finance")
    ap.add_argument("--file", help="Local FinanceData JSON. If omitted, GET /api/finance is used.")
    ap.add_argument("--scenario", choices=["current", "exist"], default=None)
    ap.add_argument("--base", default=DEFAULT_BASE)
    ap.add_argument("--secret", default="")
    ap.add_argument("--output-dir", default=None)
    ap.add_argument("--month", default=None, help="Output month label YYYY-MM")
    ap.add_argument("--format", choices=["auto", "pdf", "docx"], default="auto")
    return ap.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    if not args.secret:
        args.secret = __import__("os").environ.get("DASHBOARD_API_SECRET", "")
    try:
        data = load_or_fetch(args)
        path = build_report(
            data,
            output_dir=args.output_dir,
            output_format=args.format,
            now_label=args.month,
        )
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"OK: report={path} link={drive_link_for(path)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
